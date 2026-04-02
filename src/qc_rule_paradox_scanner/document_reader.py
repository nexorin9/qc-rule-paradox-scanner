"""规则文档读取模块

支持读取多种格式的规则文档：TXT、DOCX、PDF、Markdown

性能优化：
- 支持并发文档加载
"""

import os
from abc import ABC, abstractmethod
from concurrent.futures import ThreadPoolExecutor, as_completed
from dataclasses import dataclass
from pathlib import Path
from typing import Optional

import docx
import pdfplumber


class DocumentLoadError(Exception):
    """文档加载失败异常"""

    pass


@dataclass
class Document:
    """文档数据类"""

    content: str  # 文档文本内容
    source: str  # 文档来源路径
    doc_type: str  # 文档类型：txt, docx, pdf, markdown


class DocumentReader(ABC):
    """文档读取器基类"""

    def __init__(self, file_path: str):
        self.file_path = Path(file_path)
        if not self.file_path.exists():
            raise DocumentLoadError(f"文件不存在: {file_path}")

    @abstractmethod
    def read(self) -> Document:
        """读取文档并返回 Document 对象"""
        pass

    def _get_file_type(self) -> str:
        """根据文件扩展名确定文档类型"""
        suffix = self.file_path.suffix.lower()
        type_map = {
            ".txt": "txt",
            ".docx": "docx",
            ".pdf": "pdf",
            ".md": "markdown",
            ".markdown": "markdown",
        }
        return type_map.get(suffix, "unknown")


class TXTReader(DocumentReader):
    """纯文本文件读取器"""

    def read(self) -> Document:
        try:
            with open(self.file_path, "r", encoding="utf-8") as f:
                content = f.read()
        except UnicodeDecodeError:
            # 尝试其他编码
            for encoding in ["gbk", "gb2312", "gb18030"]:
                try:
                    with open(self.file_path, "r", encoding=encoding) as f:
                        content = f.read()
                    break
                except UnicodeDecodeError:
                    continue
            else:
                raise DocumentLoadError(f"无法解码文件 {self.file_path}，请确认文件编码为 UTF-8")

        return Document(
            content=content,
            source=str(self.file_path),
            doc_type=self._get_file_type(),
        )


class DOCXReader(DocumentReader):
    """Word 文档读取器"""

    def read(self) -> Document:
        try:
            doc = docx.Document(self.file_path)
            # 读取所有段落
            paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
            content = "\n".join(paragraphs)

            # 如果表格存在，也可以提取表格内容
            for table in doc.tables:
                for row in table.rows:
                    row_text = " | ".join(cell.text for cell in row.cells if cell.text.strip())
                    if row_text:
                        paragraphs.append(row_text)
                paragraphs.append("--- TABLE ---")

            content = "\n".join(paragraphs)

            return Document(
                content=content,
                source=str(self.file_path),
                doc_type=self._get_file_type(),
            )
        except Exception as e:
            raise DocumentLoadError(f"读取 Word 文档失败 {self.file_path}: {str(e)}")


class PDFReader(DocumentReader):
    """PDF 文档读取器"""

    def read(self) -> Document:
        try:
            with pdfplumber.open(self.file_path) as pdf:
                text_parts = []
                for page_num, page in enumerate(pdf.pages, 1):
                    text = page.extract_text()
                    if text:
                        text_parts.append(f"[页面 {page_num}]\n{text}")
                    else:
                        text_parts.append(f"[页面 {page_num}]\n[无法提取文本]")

                content = "\n\n".join(text_parts)

                return Document(
                    content=content,
                    source=str(self.file_path),
                    doc_type=self._get_file_type(),
                )
        except Exception as e:
            raise DocumentLoadError(f"读取 PDF 文档失败 {self.file_path}: {str(e)}")


class MarkdownReader(DocumentReader):
    """Markdown 文件读取器"""

    def read(self) -> Document:
        try:
            with open(self.file_path, "r", encoding="utf-8") as f:
                content = f.read()
        except UnicodeDecodeError:
            for encoding in ["gbk", "gb2312", "gb18030"]:
                try:
                    with open(self.file_path, "r", encoding=encoding) as f:
                        content = f.read()
                    break
                except UnicodeDecodeError:
                    continue
            else:
                raise DocumentLoadError(f"无法解码文件 {self.file_path}")

        return Document(
            content=content,
            source=str(self.file_path),
            doc_type=self._get_file_type(),
        )


def load_document(file_path: str) -> Document:
    """统一文档加载入口

    根据文件扩展名自动选择合适的读取器

    Args:
        file_path: 文件路径

    Returns:
        Document 对象

    Raises:
        DocumentLoadError: 当文件格式不支持或读取失败时
    """
    path = Path(file_path)
    suffix = path.suffix.lower()

    readers = {
        ".txt": TXTReader,
        ".docx": DOCXReader,
        ".pdf": PDFReader,
        ".md": MarkdownReader,
        ".markdown": MarkdownReader,
    }

    reader_class = readers.get(suffix)
    if reader_class is None:
        raise DocumentLoadError(f"不支持的文件格式: {suffix}，支持的格式: {list(readers.keys())}")

    reader = reader_class(file_path)
    return reader.read()


def load_documents(file_paths: list[str]) -> list[Document]:
    """批量加载多个文档（串行）

    Args:
        file_paths: 文件路径列表

    Returns:
        Document 对象列表
    """
    documents = []
    for file_path in file_paths:
        try:
            doc = load_document(file_path)
            documents.append(doc)
        except DocumentLoadError as e:
            print(f"警告: 跳过文件 {file_path} - {str(e)}")
    return documents


def load_documents_concurrent(
    file_paths: list[str],
    max_workers: int = 4,
    progress_callback: Optional[callable] = None,
) -> list[Document]:
    """并发批量加载多个文档

    Args:
        file_paths: 文件路径列表
        max_workers: 最大并发数
        progress_callback: 进度回调函数，签名为 callback(completed: int, total: int)

    Returns:
        Document 对象列表
    """
    documents = []
    total = len(file_paths)
    completed = 0

    with ThreadPoolExecutor(max_workers=max_workers) as executor:
        future_to_path = {
            executor.submit(load_document, fp): fp
            for fp in file_paths
        }

        for future in as_completed(future_to_path):
            completed += 1
            if progress_callback:
                progress_callback(completed, total)

            file_path = future_to_path[future]
            try:
                doc = future.result()
                documents.append(doc)
            except DocumentLoadError as e:
                print(f"警告: 跳过文件 {file_path} - {str(e)}")

    return documents


if __name__ == "__main__":
    # 简单测试
    import sys

    if len(sys.argv) > 1:
        file_path = sys.argv[1]
        try:
            doc = load_document(file_path)
            print(f"成功加载文档: {doc.source}")
            print(f"类型: {doc.doc_type}")
            print(f"内容长度: {len(doc.content)} 字符")
            print(f"前200字符:\n{doc.content[:200]}")
        except DocumentLoadError as e:
            print(f"加载失败: {e}")
    else:
        print("用法: python document_reader.py <文件路径>")